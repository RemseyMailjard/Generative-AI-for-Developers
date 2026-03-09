from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def add_learning_objectives(doc: Document, objectives: list[str]) -> None:
    doc.add_heading("Learning Objectives", level=2)
    for item in objectives:
        doc.add_paragraph(item, style="List Bullet")


def add_topic_section(
    doc: Document,
    title: str,
    why_it_matters: str,
    concepts: list[str],
    practice: list[str],
    quick_check: list[str],
) -> None:
    doc.add_heading(title, level=1)

    doc.add_paragraph("Why this matters at DocuSign", style="Intense Quote")
    doc.add_paragraph(why_it_matters)

    doc.add_heading("Core concepts", level=2)
    for item in concepts:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Hands-on practice", level=2)
    for step in practice:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Quick knowledge check", level=2)
    for question in quick_check:
        doc.add_paragraph(question, style="List Bullet")


def build_training_guide(output_path: str) -> None:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    title = doc.add_heading("DocuSign Developer Training", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Self-Learning Guide: Generative AI for Developers")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Prepared: {date.today().isoformat()}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(
        "This guide helps developers build production-ready AI features for intelligent agreement "
        "management. Each module includes concepts, practical exercises, and a quick check."
    )

    add_learning_objectives(
        doc,
        [
            "Understand core Generative AI patterns and limits in enterprise software.",
            "Write robust prompts and evaluate output quality with repeatable methods.",
            "Design RAG pipelines for trusted, source-grounded responses.",
            "Implement conversational assistants and multi-step AI agents.",
            "Integrate external capabilities safely using MCP servers and tools.",
        ],
    )

    doc.add_page_break()

    add_topic_section(
        doc,
        "1. Generative AI Fundamentals",
        "Agreement workflows require reliability and trust. Fundamentals help teams choose the "
        "right model behavior for summarization, extraction, Q&A, and content generation.",
        [
            "Foundation models, tokens, context windows, temperature, and top_p.",
            "Common model tasks: classify, summarize, transform, generate, reason.",
            "Hallucinations, latency, cost, privacy, and evaluation trade-offs.",
            "Model selection dimensions: capability, speed, cost, and governance.",
            "Responsible AI basics: PII handling, bias checks, and human-in-the-loop review.",
        ],
        [
            "Run the same prompt at temperatures 0.1, 0.5, and 0.9 and compare outputs.",
            "Classify 20 support snippets into intent labels and track error cases.",
            "Create a simple evaluation rubric: correctness, clarity, and policy compliance.",
        ],
        [
            "When would you prefer a smaller, faster model over a larger model?",
            "What are two reliable ways to reduce hallucinations?",
        ],
    )

    add_topic_section(
        doc,
        "2. Prompt Engineering",
        "Prompt quality directly impacts customer-facing outcomes such as clause explanation, "
        "contract summarization, and routing recommendations.",
        [
            "Prompt anatomy: role, objective, constraints, context, and output schema.",
            "Few-shot prompting and examples for style and format consistency.",
            "Structured outputs using JSON schemas and validation guards.",
            "Prompt iteration loop: draft, test, score, refine, and version.",
            "Prompt safety patterns: refusal boundaries and policy-aware instructions.",
        ],
        [
            "Design a prompt that extracts signer obligations into strict JSON.",
            "Add two few-shot examples to improve extraction consistency.",
            "Introduce a validation step that retries when schema checks fail.",
        ],
        [
            "Which prompt elements are mandatory for deterministic enterprise outputs?",
            "How do few-shot examples reduce formatting variability?",
        ],
    )

    add_topic_section(
        doc,
        "3. Retrieval-Augmented Generation (RAG)",
        "DocuSign use cases often require answers grounded in internal playbooks, legal guidance, "
        "and policy docs. RAG adds trusted context before generation.",
        [
            "RAG flow: ingestion, chunking, embeddings, vector search, reranking, generation.",
            "Chunking strategy and metadata design for traceability.",
            "Hybrid retrieval: keyword + vector for better recall.",
            "Grounded response design: citations, confidence, and abstain behavior.",
            "Quality metrics: hit rate, context precision, and answer faithfulness.",
        ],
        [
            "Chunk a policy document using two different chunk sizes and compare retrieval quality.",
            "Implement top-k retrieval with metadata filters (region, policy version, team).",
            "Return answers with citation links and a fallback for insufficient evidence.",
        ],
        [
            "Why is metadata critical in enterprise RAG systems?",
            "What does 'abstain' mean and when should it be used?",
        ],
    )

    add_topic_section(
        doc,
        "4. Chatbots",
        "Developers can deliver faster support and better agreement workflows with chat interfaces "
        "that handle intent, context, and handoff.",
        [
            "Conversation state management and memory boundaries.",
            "Intent routing, response templates, and escalation paths.",
            "Session context vs. persistent profile data.",
            "UX principles: transparency, confirmation, and safe fallback prompts.",
            "Operational metrics: resolution rate, CSAT, and containment.",
        ],
        [
            "Build a chatbot flow for agreement status questions and escalation to a human agent.",
            "Add response templates for identity verification and account-specific guidance.",
            "Log conversations and score outcomes against resolution KPIs.",
        ],
        [
            "When should a chatbot escalate to a human?",
            "How can you prevent context leakage across sessions?",
        ],
    )

    add_topic_section(
        doc,
        "5. AI Agents",
        "Agents can execute multi-step tasks such as gathering data, deciding actions, and calling "
        "business systems to move agreements forward.",
        [
            "Agent architecture: planner, memory, tools, and execution loop.",
            "Single-agent vs multi-agent workflows and orchestration patterns.",
            "Guardrails: action limits, approvals, and rollback strategies.",
            "Task decomposition and reliability checks for long-running actions.",
            "Observability: traces, step logs, and failure analysis.",
        ],
        [
            "Design an agent that triages inbound agreement requests and proposes next actions.",
            "Add approval checkpoints before any irreversible business operation.",
            "Simulate failure scenarios (tool timeout, bad context, invalid inputs).",
        ],
        [
            "What controls reduce risk when agents can call external systems?",
            "When is multi-agent orchestration justified over a single agent?",
        ],
    )

    add_topic_section(
        doc,
        "6. MCP Servers and Tool Integration",
        "Model Context Protocol (MCP) standardizes how AI systems discover and call tools. "
        "It enables secure, modular integrations with internal and external services.",
        [
            "MCP basics: tools, resources, prompts, and capability discovery.",
            "Designing tool contracts: clear parameters, schema validation, and errors.",
            "Security patterns: scoped credentials, least privilege, and audit logging.",
            "Timeouts, retries, circuit breakers, and idempotency for robust execution.",
            "Integration testing for tool reliability and model-tool alignment.",
        ],
        [
            "Expose a sample MCP tool that fetches agreement metadata by ID.",
            "Add schema validation and standardized error responses.",
            "Instrument requests with trace IDs and audit logs for every tool call.",
        ],
        [
            "Why should MCP tools have strict input and output schemas?",
            "What reliability patterns protect user experience during tool failures?",
        ],
    )

    doc.add_heading("Suggested 2-Week Self-Learning Plan", level=1)
    for item in [
        "Week 1: Fundamentals + Prompt Engineering + RAG fundamentals.",
        "Week 2: Chatbots + AI Agents + MCP integration capstone.",
        "Capstone: Build a grounded agreement assistant with one MCP tool.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Capstone Deliverable", level=1)
    doc.add_paragraph(
        "Build and demo a developer-focused agreement assistant that can answer grounded questions, "
        "route support intents, and invoke at least one MCP tool. Include metrics for answer quality, "
        "tool success rate, and failure handling."
    )

    doc.add_heading("Completion Checklist", level=1)
    for item in [
        "All six topic modules reviewed.",
        "At least one hands-on exercise completed per module.",
        "Capstone demo and short retrospective documented.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output_path)


if __name__ == "__main__":
    output_name = "DocuSign-Generative-AI-Developer-Training-Guide.docx"
    build_training_guide(output_name)
    print(f"Generated {output_name}")
