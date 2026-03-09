from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


def build_demo_report(output_path: str) -> None:
    doc = Document()

    # Apply explicit page setup for consistent rendering.
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    title = doc.add_heading("Quarterly Business Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Q1 2026 - Confidential")
    meta_run.font.size = Pt(12)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    summary = doc.add_paragraph()
    summary.add_run("Revenue grew ")
    summary.add_run("23% year-over-year").bold = True
    summary.add_run(", reaching $4.2M in Q1 2026 with improved customer retention.")

    doc.add_heading("Key Highlights", level=2)
    for item in [
        "Revenue: $4.2M (+23% YoY)",
        "New customers: 847",
        "Churn rate: 2.1% (down from 3.4%)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Priorities for Q2", level=2)
    for item in [
        "Launch enterprise tier",
        "Hire 3 senior engineers",
        "Expand to EU market",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Financial Summary", level=1)
    table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Metric", "Q1 2025", "Q1 2026", "Change"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data = [
        ["Revenue", "$3.4M", "$4.2M", "+23%"],
        ["Customers", "2,140", "2,987", "+40%"],
        ["MRR", "$283K", "$350K", "+24%"],
        ["Churn", "3.4%", "2.1%", "-38%"],
    ]
    for row_data in data:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Company Inc. - Confidential"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.runs[0].font.size = Pt(8)
    header_para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)


if __name__ == "__main__":
    build_demo_report("word-skill-demo-report.docx")
    print("Generated word-skill-demo-report.docx")
